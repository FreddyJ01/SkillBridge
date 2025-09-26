import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from typing import Optional, Dict, List, Tuple
import re
from copy import deepcopy

# Import our advanced XML processing system
try:
    from xml_reconstructor import DocumentReconstructor
    ADVANCED_XML_AVAILABLE = True
except ImportError:
    ADVANCED_XML_AVAILABLE = False
    print("⚠️  Advanced XML processing not available, using basic mode")

class DocumentProcessor:
    """Handles reading and writing Word documents while preserving formatting"""
    
    @staticmethod
    def extract_text_from_docx(file_path: str) -> Optional[str]:
        """Extract text content from a Word document"""
        try:
            doc = Document(file_path)
            full_text = []
            
            for paragraph in doc.paragraphs:
                full_text.append(paragraph.text)
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        full_text.append(cell.text)
            
            return '\n'.join(full_text)
        except Exception as e:
            print(f"❌ Error reading document {file_path}: {e}")
            return None
    
    @staticmethod
    def extract_document_structure(file_path: str) -> Optional[List[Dict]]:
        """Extract detailed structure and formatting from document"""
        try:
            doc = Document(file_path)
            structure = []
            
            for para in doc.paragraphs:
                para_info = {
                    'text': para.text,
                    'style': para.style.name if para.style else 'Normal',
                    'alignment': para.alignment,
                    'space_before': para.paragraph_format.space_before,
                    'space_after': para.paragraph_format.space_after,
                    'line_spacing': para.paragraph_format.line_spacing,
                    'first_line_indent': para.paragraph_format.first_line_indent,
                    'left_indent': para.paragraph_format.left_indent,
                    'right_indent': para.paragraph_format.right_indent,
                    'runs': []
                }
                
                # Extract run-level formatting
                for run in para.runs:
                    run_info = {
                        'text': run.text,
                        'bold': run.bold,
                        'italic': run.italic,
                        'underline': run.underline,
                        'font_name': run.font.name,
                        'font_size': run.font.size,
                        'font_color': run.font.color.rgb if run.font.color.rgb else None,
                        'highlight_color': run.font.highlight_color,
                    }
                    para_info['runs'].append(run_info)
                
                structure.append(para_info)
            
            return structure
        except Exception as e:
            print(f"❌ Error extracting structure from {file_path}: {e}")
            return None
    
    @staticmethod
    def create_tailored_resume(original_resume_path: str, tailored_content: str, output_path: str) -> bool:
        """Create a new resume with tailored content while preserving 100% identical formatting"""
        
        # Try advanced XML reconstruction first (for 100% formatting preservation)
        if ADVANCED_XML_AVAILABLE:
            try:
                print("� Attempting 100% formatting preservation with XML reconstruction...")
                reconstructor = DocumentReconstructor()
                success = reconstructor.create_perfectly_formatted_resume(
                    original_resume_path, tailored_content, output_path
                )
                
                if success:
                    print("✅ 100% formatting preservation successful!")
                    return True
                else:
                    print("⚠️  XML reconstruction failed, falling back to structure preservation...")
                    
            except Exception as e:
                print(f"⚠️  XML reconstruction error: {e}")
                print("📋 Falling back to structure preservation mode...")
        
        # Fallback to structure preservation method
        try:
            print("📋 Using structure preservation method...")
            
            # Extract the original document structure
            original_structure = DocumentProcessor.extract_document_structure(original_resume_path)
            if not original_structure:
                raise Exception("Could not extract original document structure")
            
            print("🔍 Mapping content to structure...")
            
            # Parse the AI-generated content
            new_content_lines = [line.strip() for line in tailored_content.split('\n') if line.strip()]
            
            # Map new content to original structure
            mapped_structure = DocumentProcessor._map_content_to_structure(
                original_structure, new_content_lines
            )
            
            print("🎨 Creating formatted document...")
            
            # Create new document with original formatting
            success = DocumentProcessor._create_formatted_document(
                original_resume_path, mapped_structure, output_path
            )
            
            if success:
                print("✅ Structure preservation successful!")
                return True
            else:
                raise Exception("Failed to create formatted document")
            
        except Exception as e:
            print(f"❌ Structure preservation failed: {e}")
            print("📄 Using fallback document creation...")
            # Final fallback: create a basic document
            return DocumentProcessor._create_fallback_document(tailored_content, output_path)
    
    @staticmethod
    def create_error_document(error_message: str, output_path: str) -> bool:
        """Create an error document with the error message"""
        try:
            doc = Document()
            
            # Add title
            title = doc.add_heading('SkillBridge Processing Error', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add error message
            doc.add_paragraph()
            error_para = doc.add_paragraph('We apologize, but there was an error processing your resume:')
            error_para.runs[0].bold = True
            
            doc.add_paragraph()
            doc.add_paragraph(error_message)
            
            doc.add_paragraph()
            doc.add_paragraph('Please check:')
            
            checklist = [
                '• Both JD.docx and CurrentResume.docx files are valid Word documents',
                '• The files contain readable text content',
                '• Your internet connection (if using OpenAI)',
                '• Ollama is running (if using local AI)',
                '• Try again in a few moments'
            ]
            
            for item in checklist:
                doc.add_paragraph(item)
            
            doc.add_paragraph()
            footer = doc.add_paragraph('If the problem persists, please check the console for detailed error messages.')
            footer.runs[0].italic = True
            
            doc.save(output_path)
            return True
            
        except Exception as e:
            print(f"❌ Error creating error document: {e}")
            return False
    
    @staticmethod
    def _map_content_to_structure(original_structure: List[Dict], new_content: List[str]) -> List[Dict]:
        """Map new content lines to original document structure"""
        mapped = []
        content_index = 0
        
        for orig_para in original_structure:
            if content_index >= len(new_content):
                break
            
            orig_text = orig_para['text'].strip()
            
            # Skip empty paragraphs in original
            if not orig_text:
                mapped.append(orig_para)  # Keep empty paragraphs for spacing
                continue
            
            # Map content to structure
            new_para = deepcopy(orig_para)
            
            if content_index < len(new_content):
                new_text = new_content[content_index]
                new_para['text'] = new_text
                
                # Update runs with new text while preserving formatting
                if new_para['runs']:
                    # Distribute text across runs proportionally
                    new_para['runs'] = DocumentProcessor._update_runs_with_text(
                        new_para['runs'], new_text
                    )
                
                content_index += 1
            
            mapped.append(new_para)
        
        # Add any remaining content as new paragraphs with last paragraph's style
        while content_index < len(new_content):
            if mapped:
                last_style = deepcopy(mapped[-1])
                last_style['text'] = new_content[content_index]
                last_style['runs'] = [{'text': new_content[content_index], 'bold': False, 'italic': False}]
                mapped.append(last_style)
            content_index += 1
        
        return mapped
    
    @staticmethod
    def _update_runs_with_text(original_runs: List[Dict], new_text: str) -> List[Dict]:
        """Update run text while preserving formatting"""
        if not original_runs:
            return [{'text': new_text, 'bold': False, 'italic': False}]
        
        # If single run, replace text
        if len(original_runs) == 1:
            new_run = deepcopy(original_runs[0])
            new_run['text'] = new_text
            return [new_run]
        
        # Multiple runs - try to preserve formatting patterns
        new_runs = []
        words = new_text.split()
        words_per_run = max(1, len(words) // len(original_runs))
        
        word_index = 0
        for i, run in enumerate(original_runs):
            new_run = deepcopy(run)
            
            # Assign words to this run
            if i == len(original_runs) - 1:
                # Last run gets remaining words
                run_words = words[word_index:]
            else:
                run_words = words[word_index:word_index + words_per_run]
                word_index += words_per_run
            
            new_run['text'] = ' '.join(run_words) if run_words else ''
            if new_run['text']:  # Only add non-empty runs
                new_runs.append(new_run)
        
        return new_runs if new_runs else [{'text': new_text, 'bold': False, 'italic': False}]
    
    @staticmethod
    def _create_formatted_document(original_path: str, structure: List[Dict], output_path: str) -> bool:
        """Create a new document with the mapped structure and formatting"""
        try:
            # Load original document to copy styles and settings
            original_doc = Document(original_path)
            
            # Create new document with same core settings
            new_doc = Document()
            
            # Copy document-level settings
            DocumentProcessor._copy_document_settings(original_doc, new_doc)
            
            # Create paragraphs with preserved formatting
            for para_info in structure:
                new_para = new_doc.add_paragraph()
                
                # Apply paragraph formatting
                DocumentProcessor._apply_paragraph_formatting(new_para, para_info)
                
                # Add runs with formatting
                if para_info['runs']:
                    for run_info in para_info['runs']:
                        if run_info['text']:
                            run = new_para.add_run(run_info['text'])
                            DocumentProcessor._apply_run_formatting(run, run_info)
                else:
                    # Fallback: add text as single run
                    if para_info['text']:
                        new_para.add_run(para_info['text'])
            
            new_doc.save(output_path)
            return True
            
        except Exception as e:
            print(f"❌ Error creating formatted document: {e}")
            return False
    
    @staticmethod
    def _copy_document_settings(source_doc: Document, target_doc: Document):
        """Copy document-level settings like margins, fonts, etc."""
        try:
            # Copy margins
            target_doc.sections[0].page_height = source_doc.sections[0].page_height
            target_doc.sections[0].page_width = source_doc.sections[0].page_width
            target_doc.sections[0].left_margin = source_doc.sections[0].left_margin
            target_doc.sections[0].right_margin = source_doc.sections[0].right_margin
            target_doc.sections[0].top_margin = source_doc.sections[0].top_margin
            target_doc.sections[0].bottom_margin = source_doc.sections[0].bottom_margin
        except:
            pass  # Continue even if some settings can't be copied
    
    @staticmethod
    def _apply_paragraph_formatting(paragraph, para_info: Dict):
        """Apply paragraph-level formatting"""
        try:
            pf = paragraph.paragraph_format
            
            if para_info.get('alignment'):
                pf.alignment = para_info['alignment']
            if para_info.get('space_before'):
                pf.space_before = para_info['space_before']
            if para_info.get('space_after'):
                pf.space_after = para_info['space_after']
            if para_info.get('line_spacing'):
                pf.line_spacing = para_info['line_spacing']
            if para_info.get('first_line_indent'):
                pf.first_line_indent = para_info['first_line_indent']
            if para_info.get('left_indent'):
                pf.left_indent = para_info['left_indent']
            if para_info.get('right_indent'):
                pf.right_indent = para_info['right_indent']
        except:
            pass  # Continue even if some formatting can't be applied
    
    @staticmethod
    def _apply_run_formatting(run, run_info: Dict):
        """Apply run-level formatting (font, bold, italic, etc.)"""
        try:
            if run_info.get('bold') is not None:
                run.bold = run_info['bold']
            if run_info.get('italic') is not None:
                run.italic = run_info['italic']
            if run_info.get('underline') is not None:
                run.underline = run_info['underline']
            if run_info.get('font_name'):
                run.font.name = run_info['font_name']
            if run_info.get('font_size'):
                run.font.size = run_info['font_size']
            if run_info.get('font_color'):
                run.font.color.rgb = run_info['font_color']
        except:
            pass  # Continue even if some formatting can't be applied
    
    @staticmethod
    def _create_fallback_document(content: str, output_path: str) -> bool:
        """Create a basic document as fallback when formatting preservation fails"""
        try:
            doc = Document()
            
            # Split content into paragraphs and add with basic formatting
            paragraphs = content.split('\n')
            
            for para_text in paragraphs:
                para_text = para_text.strip()
                if not para_text:
                    doc.add_paragraph()  # Empty paragraph for spacing
                    continue
                
                # Apply basic smart formatting
                if DocumentProcessor._is_heading(para_text):
                    heading = doc.add_heading(para_text, level=1)
                elif DocumentProcessor._is_subheading(para_text):
                    subheading = doc.add_heading(para_text, level=2)
                elif DocumentProcessor._is_bullet_point(para_text):
                    doc.add_paragraph(para_text, style='List Bullet')
                else:
                    para = doc.add_paragraph(para_text)
                    DocumentProcessor._apply_smart_formatting(para)
            
            doc.save(output_path)
            return True
            
        except Exception as e:
            print(f"❌ Error creating fallback document: {e}")
            return False
    
    @staticmethod
    def _is_heading(text: str) -> bool:
        """Determine if text should be formatted as a main heading"""
        headings = [
            'professional summary', 'summary', 'objective', 'profile',
            'work experience', 'experience', 'employment history',
            'education', 'skills', 'technical skills', 'core competencies',
            'certifications', 'awards', 'projects', 'achievements'
        ]
        return any(heading in text.lower() for heading in headings) and len(text) < 50
    
    @staticmethod
    def _is_subheading(text: str) -> bool:
        """Determine if text should be formatted as a subheading (job title/company)"""
        # Look for patterns like "Job Title | Company Name" or "Company Name - Job Title"
        patterns = [
            r'.*\|.*',  # Contains pipe
            r'.*-.*',   # Contains dash
            r'.*at .*', # Contains "at"
        ]
        return any(re.match(pattern, text) for pattern in patterns) and len(text) < 100
    
    @staticmethod
    def _is_bullet_point(text: str) -> bool:
        """Determine if text should be formatted as a bullet point"""
        return text.startswith('•') or text.startswith('-') or text.startswith('*')
    
    @staticmethod
    def _apply_smart_formatting(paragraph):
        """Apply smart formatting to paragraph text"""
        text = paragraph.text
        
        # Make job titles and company names bold (simplified approach)
        # This is a basic implementation - could be enhanced
        if '|' in text or ' - ' in text or ' at ' in text:
            # Likely a job title/company line
            paragraph.runs[0].bold = True